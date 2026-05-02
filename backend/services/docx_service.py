"""
services/docx_service.py — Extract text from DOCX files using python-docx
"""
import io
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file given its bytes.
    Returns concatenated text from all paragraphs.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")

    if not paragraphs:
        raise ValueError("No text could be extracted from this DOCX file.")

    return "\n\n".join(paragraphs)
